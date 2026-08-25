# -*- coding: utf-8 -*-
"""
生成《第三帧事务所 - 素材导入与网站管理教程》Word 文档
使用：python3 scripts/generate_tutorial_docx.py
输出：项目根目录下的 "第三帧事务所_素材导入与网站管理教程.docx"
"""
import os
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============ 路径 ============
PROJECT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_FILE = os.path.join(PROJECT_DIR, "第三帧事务所_素材导入与网站管理教程.docx")


# ============ 字体与样式工具 ============
# macOS 默认中文字体：苹方（PingFang SC）优先，其次是 Heiti SC / Songti SC
# Windows 下常见 Microsoft YaHei / SimSun
def set_run_font(run, cn_name="PingFang SC", en_name="Helvetica Neue", size=11, bold=False, color=None):
    """设置 run 的中英文字体、字号、粗体、颜色"""
    run.font.name = en_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # 设置东亚字体（中文）
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), cn_name)
    rFonts.set(qn("w:ascii"), en_name)
    rFonts.set(qn("w:hAnsi"), en_name)


def set_cell_font(cell, cn_name="PingFang SC", en_name="Helvetica Neue", size=10.5, bold=False, color=None, align=None):
    """一次性设置整个 cell 里所有段落的字体"""
    for p in cell.paragraphs:
        if align:
            p.alignment = align
        for run in p.runs:
            set_run_font(run, cn_name, en_name, size, bold, color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(doc, text, level=1):
    """添加自定义样式的标题（中文友好）"""
    h = doc.add_heading("", level=level)
    run = h.add_run(text)
    cn_map = {1: "PingFang SC", 2: "PingFang SC", 3: "PingFang SC", 4: "PingFang SC"}
    sz_map = {1: 22, 2: 18, 3: 15, 4: 13}
    color_map = {
        1: RGBColor(0x1D, 0x1D, 0x1F),
        2: RGBColor(0x33, 0x33, 0x38),
        3: RGBColor(0x3A, 0x3A, 0x3C),
        4: RGBColor(0x55, 0x55, 0x55),
    }
    set_run_font(run, cn_map.get(level, "PingFang SC"), "Helvetica Neue", sz_map.get(level, 12), bold=True,
                 color=color_map.get(level))
    # 标题前后间距
    pf = h.paragraph_format
    if level == 1:
        pf.space_before = Pt(28)
        pf.space_after = Pt(14)
    elif level == 2:
        pf.space_before = Pt(22)
        pf.space_after = Pt(10)
    else:
        pf.space_before = Pt(14)
        pf.space_after = Pt(6)
    # 一级标题下方加一条细线
    if level == 1:
        pPr = h._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "C7C7CC")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return h


def add_paragraph(doc, text, size=11, bold=False, color=None, align=None, indent=None, cn="PingFang SC", en="Helvetica Neue"):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = indent
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, cn, en, size, bold, color)
    return p


def add_bullet(doc, text, level=0, size=11):
    """添加带缩进的项目符号（手工模拟，避免跨平台兼容问题）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.6)
    p.paragraph_format.space_after = Pt(3)
    bullet_char = "●  " if level == 0 else "○  " if level == 1 else "▸  "
    run = p.add_run(bullet_char + text)
    set_run_font(run, size=size)
    return p


def add_number(doc, num, text, size=11):
    """添加编号项"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{num}.  ")
    set_run_font(run, size=size, bold=True, color=RGBColor(0x00, 0x71, 0xE5))
    run2 = p.add_run(text)
    set_run_font(run2, size=size)
    return p


def add_code_block(doc, code_text, size=10):
    """添加代码/终端命令样式块（灰色背景+等宽字体）"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.6)
    pf.right_indent = Cm(0.6)
    pf.space_before = Pt(4)
    pf.space_after = Pt(6)
    # 设置段落背景灰
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F7")
    pPr.append(shd)
    # 左侧竖线
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "0071E5")
    pBdr.append(left)
    pPr.append(pBdr)

    run = p.add_run(code_text)
    set_run_font(run, cn_name="PingFang SC", en_name="Menlo", size=size, color=RGBColor(0x1D, 0x1D, 0x1F))
    return p


def add_tip(doc, tip_text, kind="info"):
    """添加提示 / 警告 / 小贴士框"""
    kind_styles = {
        "info":    {"icon": "ℹ️", "fill": "E8F0FE", "line": "0071E5", "title": "提示"},
        "warn":    {"icon": "⚠️", "fill": "FFF7E0", "line": "FF9500", "title": "注意"},
        "danger":  {"icon": "⛔", "fill": "FFE5E5", "line": "FF3B30", "title": "警告"},
        "success": {"icon": "✅", "fill": "E5F7EC", "line": "34C759", "title": "成功"},
    }
    s = kind_styles.get(kind, kind_styles["info"])
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.4)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), s["fill"])
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), s["line"])
    pBdr.append(left)
    pPr.append(pBdr)
    r1 = p.add_run(f"{s['icon']}  {s['title']}：")
    set_run_font(r1, size=10.5, bold=True, color=RGBColor(0x1D, 0x1D, 0x1F))
    r2 = p.add_run(tip_text)
    set_run_font(r2, size=10.5, color=RGBColor(0x3A, 0x3A, 0x3C))
    return p


def add_table(doc, headers, rows, col_widths_cm=None, first_col_bold=False):
    """添加样式化的表格"""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 表头
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        # 背景色
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1D1D1F")
        tcPr.append(shd)
        # 字体
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 内容行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = str(val)
            if i % 2 == 1:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F5F5F7")
                tcPr.append(shd)
            bold = first_col_bold and j == 0
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10, bold=bold)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 列宽
    if col_widths_cm:
        for row in t.rows:
            for j, w in enumerate(col_widths_cm):
                row.cells[j].width = Cm(w)
    doc.add_paragraph()  # 加个空行
    return t


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "dotted")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C7C7CC")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============ 主流程 ============
def build_document():
    doc = Document()

    # 默认边距：2cm 左右上下
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    # ---------------------
    # 封面
    # ---------------------
    for _ in range(4):
        doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("Third Frame Co")
    set_run_font(r, "PingFang SC", "Helvetica Neue", 36, bold=True, color=RGBColor(0x1D, 0x1D, 0x1F))

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run("第 三 帧 事 务 所")
    set_run_font(r2, "Songti SC", "Helvetica Neue", 20, color=RGBColor(0x63, 0x63, 0x66))

    doc.add_paragraph()
    doc.add_paragraph()

    mid_p = doc.add_paragraph()
    mid_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = mid_p.add_run("素材导入与网站管理教程")
    set_run_font(r3, "PingFang SC", "Helvetica Neue", 24, bold=True, color=RGBColor(0x00, 0x71, 0xE5))

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = date_p.add_run(f"版本 v1.0    生成日期 {datetime.now().strftime('%Y-%m-%d')}")
    set_run_font(r4, size=11, color=RGBColor(0x8E, 0x8E, 0x93))

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = info_p.add_run("免费可商用 Log 素材库 · 专注调色练习")
    set_run_font(r5, size=11, color=RGBColor(0x8E, 0x8E, 0x93))

    # 分页
    doc.add_page_break()

    # ---------------------
    # 目录
    # ---------------------
    add_heading(doc, "目  录", level=1)
    toc_items = [
        ("01", "欢迎与网站简介"),
        ("02", "关键信息速查表"),
        ("03", "日常工作流程总览"),
        ("04", "第一步：登录后台管理面板"),
        ("05", "第二步：管理多级分类（品牌 / Log / 场景）"),
        ("06", "第三步：录入新素材（字段全解）"),
        ("07", "第四步：准备封面图与预览视频"),
        ("08", "第五步：导出数据并替换本地文件"),
        ("09", "第六步：一键部署到 GitHub Pages"),
        ("10", "进阶：修改关于我 & 首页精选"),
        ("11", "进阶：修改后台密码 / 站点名称"),
        ("12", "常见问题 FAQ"),
        ("附录 A", "字段速查表"),
        ("附录 B", "项目目录结构"),
        ("附录 C", "重新生成 GitHub Personal Access Token"),
    ]
    for num, name in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1)
        r1 = p.add_run(f"{num}  ")
        set_run_font(r1, size=12, bold=True, color=RGBColor(0x00, 0x71, 0xE5))
        r2 = p.add_run(name)
        set_run_font(r2, size=12)

    doc.add_page_break()

    # ---------------------
    # 01 欢迎与简介
    # ---------------------
    add_heading(doc, "01  欢迎与网站简介", level=1)
    add_paragraph(doc, "欢迎使用「第三帧事务所」素材共享站点！这是一个基于 GitHub Pages 零成本部署的纯静态网站，"
                       "专注于分享你个人拍摄的免费可商用 Log 灰片、照片素材，供调色练习与创作使用。", size=11)
    add_paragraph(doc, "本网站具备以下能力：")
    add_bullet(doc, "多维度筛选：相机品牌、Log 灰度、照片/视频、室内/室外、分辨率（4K / 1080P）、帧率等")
    add_bullet(doc, "中文模糊搜索：素材标题、描述、标签内任意关键词")
    add_bullet(doc, "明暗主题切换：参考 Apple 官网风格的浅色/深色界面")
    add_bullet(doc, "低分辨率小样在线预览 + 夸克/百度网盘原片跳转下载")
    add_bullet(doc, "密码保护的后台管理面板 + 多级分类自定义（最多 3 级）")
    add_bullet(doc, "草稿本地保存 + 一键导出 JSON，纯前端无需数据库")

    add_tip(doc, "本教程针对非技术背景用户编写，所有命令与操作均可复制粘贴使用。如遇问题，请翻到最后一章常见问题 FAQ。", "info")

    # ---------------------
    # 02 关键信息速查表
    # ---------------------
    add_heading(doc, "02  关键信息速查表", level=1)
    add_table(doc,
              ["项目", "值 / 地址", "说明"],
              [
                  ["🌐 网站公开地址", "https://z1sj.github.io/third-frame-co/", "任何人都可以访问"],
                  ["📂 GitHub 代码仓库", "https://github.com/z1sj/third-frame-co", "源代码、提交历史"],
                  ["🔧 后台管理入口", "网站右上角「管理员」按钮 / #/admin", "输入密码进入"],
                  ["🔑 后台默认密码", "thirdframe2026", "建议立即修改（见第 11 章）"],
                  ["💻 本地预览地址", "http://localhost:5173/", "开发服务器运行时可查看"],
                  ["📁 项目根目录（本机）", PROJECT_DIR, "终端 cd 到此处执行命令"],
                  ["📦 素材 JSON 数据", "./src/data/materials.json", "从后台导出后覆盖此文件"],
                  ["🗂️  分类 JSON 数据", "./src/data/categories.json", "从后台导出后覆盖此文件"],
                  ["🖼️  封面图目录", "./public/covers/", "命名：<素材ID>.jpg"],
                  ["🎞️  预览视频目录", "./public/previews/", "命名：<素材ID>.mp4"],
                  ["🚀 一键部署命令", "./deploy.sh", "构建 + 发布到 GitHub Pages"],
              ],
              col_widths_cm=[3.8, 7.2, 5.0],
              first_col_bold=True)

    # ---------------------
    # 03 日常工作流程总览
    # ---------------------
    add_heading(doc, "03  日常工作流程总览", level=1)
    add_paragraph(doc, "每次你拍摄了新的素材，需要把它发布到网站上，请严格按下面 6 步顺序操作。只要按顺序来，永远不会出错：")
    doc.add_paragraph()

    steps = [
        ("① 登录后台", "打开网站 → 点右上角管理员 → 输入密码 thirdframe2026 进入后台"),
        ("② 管理分类", "先确认分类是否需要新增（比如新买的新相机、新的 Log 类型），不需要则跳过"),
        ("③ 录入素材", "点「新增素材」，填好标题、分类、夸克链接、参数等全部字段，保存草稿"),
        ("④ 放预览文件", "在本机把封面图复制到 public/covers/<ID>.jpg，预览视频复制到 public/previews/<ID>.mp4"),
        ("⑤ 导出 JSON", "后台点「导出全部数据」，把解压后的 materials.json 与 categories.json 覆盖到 src/data/"),
        ("⑥ 一键部署", "终端 cd 到项目根目录，执行 ./deploy.sh，等待 1-3 分钟后网站自动更新"),
    ]
    for title, desc in steps:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(6)
        r1 = p.add_run(f"{title}    ")
        set_run_font(r1, size=12, bold=True, color=RGBColor(0x00, 0x71, 0xE5))
        r2 = p.add_run(desc)
        set_run_font(r2, size=11)

    add_tip(doc, "小贴士：素材录入和预览文件的准备没有严格的先后顺序，可以交叉进行。关键是「导出 JSON」要在「执行部署」之前完成。", "info")

    # ---------------------
    # 04 登录后台
    # ---------------------
    add_heading(doc, "04  第一步：登录后台管理面板", level=1)
    add_heading(doc, "4.1  进入后台的两种方式", level=2)
    add_number(doc, 1, "从网站右上角进入：任意页面点导航栏最右侧的「管理员」按钮 → 弹出密码输入框 → 输入密码 → 自动跳转到后台")
    add_number(doc, 2, "直接访问：在浏览器地址栏输入 网站地址/#/admin （本地预览就是 http://localhost:5173/#/admin）")

    add_heading(doc, "4.2  忘记密码怎么办", level=2)
    add_paragraph(doc, "打开项目文件 src/config.ts，找到 ADMIN_PASSWORD 字段，引号里的内容就是当前密码。修改保存后需要重新执行一次 ./deploy.sh 才会生效。（详见第 11 章）")

    # ---------------------
    # 05 管理多级分类
    # ---------------------
    add_heading(doc, "05  第二步：管理多级分类", level=1)
    add_paragraph(doc, "后台的「分类管理」Tab 可以自定义最多 3 级分类树。推荐按下面的维度组织分类：")
    add_table(doc,
              ["层级", "分类示例", "说明"],
              [
                  ["一级", "相机品牌（索尼 / 佳能 / 尼康 / 富士…）", "最顶层的分类"],
                  ["二级", "Sony / A7M4, Canon / R5…", "具体的机型或系列"],
                  ["三级", "S-Log3 / S-Cinetone / C-Log3…", "具体的 Log 灰度或色彩模式"],
              ],
              col_widths_cm=[2, 6.5, 7.5], first_col_bold=True)

    add_heading(doc, "5.1  分类操作说明", level=2)
    add_number(doc, 1, "新增一级分类：点击「+ 新增一级」→ 填入名称 → 回车保存")
    add_number(doc, 2, "新增二级/三级分类：把鼠标移到父分类上 → 点旁边的「+ 子分类」→ 填入名称")
    add_number(doc, 3, "重命名：把鼠标移到分类名上 → 点「编辑」→ 改名后回车")
    add_number(doc, 4, "删除分类：点分类旁边的「删除」按钮，该分类下所有子分类会一并删除")
    add_tip(doc, "删除分类之前，先确认没有素材还在使用这条分类路径。如果有，请先把对应素材的分类改掉，否则筛选时会看不到该素材。", "warn")

    # ---------------------
    # 06 录入新素材
    # ---------------------
    add_heading(doc, "06  第三步：录入新素材（字段全解）", level=1)
    add_paragraph(doc, "后台「素材管理」Tab，点「+ 新增素材」打开录入面板。下面是每个字段的详细说明：")

    field_rows = [
        ["素材 ID", "必填，全局唯一。建议规则：相机前缀_日期_序号，如 SONY_20260824_01",
         "⚠️ 极其重要：后面封面图、预览视频文件的命名都必须和这个 ID 完全一致。建议字母大写，用下划线分隔，不要用中文或空格。"],
        ["素材标题", "必填，简洁描述拍摄对象，如「逆光海边日落 S-Log3」", "会显示在卡片和列表里，建议控制在 30 字以内。"],
        ["素材类型", "video（视频）/ photo（照片）", "二选一，决定是否显示时长和预览视频播放按钮。"],
        ["多级分类", "必填，最多选 3 级：一级 → 二级 → 三级", "例如：相机品牌 → 索尼 → S-Log3，用于筛选。"],
        ["分辨率", "4K / 1080P / 其他自定义字符串", "影响「分辨率筛选器」的匹配，尽量统一 4K 或 1080P。"],
        ["帧率", "整数，比如 25 / 30 / 60 / 120", "影响「帧率筛选器」的匹配。照片可留空。"],
        ["拍摄场景", "室内 / 室外 / 不限", "影响「场景筛选器」的匹配。"],
        ["视频时长", "仅视频填，格式 mm:ss，如 00:42", "显示在素材卡片的右下角。"],
        ["夸克网盘链接", "必填！你发布到夸克网盘的分享链接", "点击下载按钮的主要跳转地址。建议提前测试链接是否有效。"],
        ["百度网盘链接", "可选，夸克打不开时的备用链接", "有的话就填，留空也没关系。"],
        ["封面文件名", "必填，如 SONY_20260824_01.jpg", "⚠️ 必须和你放到 public/covers/ 里的文件名（含扩展名）完全一致。"],
        ["预览视频文件名", "视频必填；照片留空，如 SONY_20260824_01.mp4", "⚠️ 必须和你放到 public/previews/ 里的文件名完全一致。建议 720P 左右的低码率压缩小样。"],
        ["素材描述", "拍摄思路、使用建议、调色思路等任意文字，支持多段", "会显示在素材详情页里，越长越好。"],
        ["标签 Tags", "多个关键词用逗号或回车分隔，如 逆光,海边,日落,人像", "搜索范围包含标签，建议多填便于搜索命中。"],
        ["相机参数（可选）", "机型、镜头、光圈、快门、ISO、白平衡", "展示在详情页的「拍摄参数」区，追求完整的话建议填。"],
        ["首页精选", "打勾 / 不打勾", "打勾后会出现在首页的精选推荐栏里。建议只给最精品的素材打勾。"],
    ]
    add_table(doc, ["字段名", "填写要求", "注意事项 / 说明"], field_rows, col_widths_cm=[3.2, 6.0, 6.8])

    add_heading(doc, "6.1  录入完成后保存草稿", level=2)
    add_paragraph(doc, "点保存后，数据会存到你当前浏览器的 localStorage。关闭网页也不会丢。但是！")
    add_tip(doc, "浏览器清理缓存、换浏览器、换电脑都会丢失草稿！所以一定要导出 JSON 备份（第 08 章）。后台每隔 30 秒会自动保存一次草稿，但这不等同于导出备份。", "warn")

    # ---------------------
    # 07 准备封面图和预览视频
    # ---------------------
    add_heading(doc, "07  第四步：准备封面图与预览视频", level=1)
    add_heading(doc, "7.1  文件命名与存放位置", level=2)
    add_paragraph(doc, "假设你的素材 ID 是 SONY_20260824_01，那么两个文件必须严格按下面的规则放：")
    add_table(doc,
              ["文件类型", "保存路径（相对于项目根目录）", "文件名（必须等于素材ID + 对应扩展名）", "推荐规格"],
              [
                  ["封面图", "public/covers/", "SONY_20260824_01.jpg", "16:9 横版，1920×1080 JPG，画质 80%，体积 < 500KB"],
                  ["预览视频", "public/previews/", "SONY_20260824_01.mp4", "1280×720 或 1600×900，码率 2-4 Mbps，时长 10-30 秒，H.264 + AAC，体积 < 20MB"],
              ],
              col_widths_cm=[2.0, 5.0, 5.0, 4.0], first_col_bold=True)
    add_tip(doc, "文件命名错误是最常见的问题。请反复确认：素材 ID → 封面文件名 → 预览文件名，三者的前缀要一模一样，大小写敏感。", "danger")

    add_heading(doc, "7.2  导出低码率预览视频（推荐命令）", level=2)
    add_paragraph(doc, "使用 FFmpeg 一行命令把原片压成 720P 小体积预览版（如果电脑没有 FFmpeg，也可以在剪映里 720P 导出）：")
    add_code_block(doc,
                   "ffmpeg -i 原片.mov -vf scale=1280:-2 -c:v libx264 -preset medium -crf 28 -c:a aac -b:a 96k -movflags +faststart SONY_20260824_01.mp4")

    # ---------------------
    # 08 导出数据
    # ---------------------
    add_heading(doc, "08  第五步：导出数据并替换本地文件", level=1)
    add_heading(doc, "8.1  导出流程", level=2)
    add_number(doc, 1, "后台点右侧「导出全部数据」按钮（注意是导出「全部」，不是单个素材）")
    add_number(doc, 2, "浏览器会下载一个压缩包，通常叫 third-frame-data-日期.zip")
    add_number(doc, 3, "解压这个 zip，你会得到两个文件：materials.json 和 categories.json")
    add_number(doc, 4, "把这两个文件复制 → 粘贴到项目目录的 src/data/ 文件夹，覆盖原文件")

    add_heading(doc, "8.2  本地替换验证（可选）", level=2)
    add_paragraph(doc, "如果你启动了本地预览服务器（http://localhost:5173/），刷新一下浏览器应该能立刻看到新的素材出现在网站里。如果看不到，先按下面清单检查：")
    add_bullet(doc, "文件名是不是 materials.json / categories.json，不是别的（比如 素材.json）")
    add_bullet(doc, "有没有放到 src/data/ 而不是项目根目录或其他目录")
    add_bullet(doc, "是不是真的覆盖了（看文件的修改日期是不是刚才的时间）")
    add_bullet(doc, "浏览器是不是强制刷新了（Command/Control + Shift + R 清空缓存刷新）")

    # ---------------------
    # 09 一键部署
    # ---------------------
    add_heading(doc, "09  第六步：一键部署到 GitHub Pages", level=1)
    add_paragraph(doc, "这一步会把代码打包，并自动发布到 GitHub 服务器上，让全世界的访客都能看到更新后的网站。")

    add_heading(doc, "9.1  执行部署（2 条命令）", level=2)
    add_paragraph(doc, "打开「终端」（macOS 启动台 → 其他 → 终端），依次执行：")
    add_code_block(doc,
                   "cd \"/Users/genhaosan/Library/Application Support/TRAE SOLO CN/ModularData/ai-agent/work-mode-projects/6a8bb8f8fe83435acf288eb9\"\n"
                   "./deploy.sh")
    add_paragraph(doc, "执行过程你会看到类似如下输出：")
    add_code_block(doc,
                   "🔨 开始构建生产版本...\n"
                   "vite v5.x building for production...\n"
                   "✓ 1924 modules transformed.\n"
                   "✓ built in 1.33s\n"
                   "📦 正在发布到 gh-pages 分支...\n"
                   "Published\n"
                   "✅ 发布成功！")

    add_heading(doc, "9.2  首次 push 或新电脑首次部署", level=2)
    add_paragraph(doc, "如果终端提示你输入 GitHub 用户名和密码：")
    add_bullet(doc, "用户名：z1sj")
    add_bullet(doc, "密码：不是你的 GitHub 登录密码！而是填你之前生成的 Personal Access Token（ghp_ 开头那串）。如果忘记了，按附录 C 重新生成一个。")
    add_tip(doc, "输入 Token 时终端不会显示星号或字符，这是正常的，粘贴完直接回车即可。", "info")

    add_heading(doc, "9.3  等待生效", level=2)
    add_paragraph(doc, "出现 Published 字样后，表示代码已推送成功。GitHub Pages 需要 1-3 分钟同步，最多 10 分钟，这段时间你刷新看到的是旧版本或 404 都属正常。")

    # ---------------------
    # 10 进阶：关于我 & 首页精选
    # ---------------------
    add_heading(doc, "10  进阶：修改关于我 & 首页精选", level=1)
    add_heading(doc, "10.1  修改「关于我」页面", level=2)
    add_paragraph(doc, "后台点「关于我」Tab，可以修改：")
    add_bullet(doc, "昵称 + 个人简介")
    add_bullet(doc, "头像（URL 或 base64 内嵌）")
    add_bullet(doc, "社交链接：名称 + 链接（小红书 / B站 / 邮箱 / 微博等任意）")
    add_paragraph(doc, "改完后记得走一遍「导出 JSON → 覆盖文件 → ./deploy.sh」的完整流程，否则不会上线。")

    add_heading(doc, "10.2  设置首页精选素材", level=2)
    add_paragraph(doc, "首页会展示 8 个（可配置）「精选」素材。两种方式设置：")
    add_number(doc, 1, "方式 A：编辑素材时勾选「首页精选」复选框（推荐，最简单）")
    add_number(doc, 2, "方式 B：在「站点设置」Tab 里手动填入 8 个精选素材的 ID")

    # ---------------------
    # 11 进阶：修改密码
    # ---------------------
    add_heading(doc, "11  进阶：修改后台密码 / 站点名称", level=1)
    add_heading(doc, "11.1  修改后台登录密码", level=2)
    add_paragraph(doc, "用任意代码编辑器（比如 VS Code、Sublime Text，甚至「文本编辑」）打开项目里的这个文件：")
    add_code_block(doc, "src/config.ts")
    add_paragraph(doc, "找到这一行：")
    add_code_block(doc, "export const ADMIN_PASSWORD = 'thirdframe2026'")
    add_paragraph(doc, "把引号里的 thirdframe2026 改成你自己的密码（建议 8 位以上，字母+数字混合，别用特殊字符），保存。")

    add_heading(doc, "11.2  修改网站名称 / Slogan", level=2)
    add_paragraph(doc, "同一个文件里：")
    add_code_block(doc,
                   "export const SITE_NAME = 'Third Frame Co'\n"
                   "export const SITE_NAME_CN = '第三帧事务所'\n"
                   "export const SITE_SLOGAN = '免费可商用素材库 · 专注 Log 调色练习'")
    add_paragraph(doc, "按需修改后保存。所有修改完成后，必须执行 ./deploy.sh 重新部署才会生效。")
    add_tip(doc, "所有 src/config.ts 里的修改：改完立即生效是在本地开发服务器（npm run dev），生产部署必须重新 ./deploy.sh。", "warn")

    # ---------------------
    # 12 常见问题 FAQ
    # ---------------------
    add_heading(doc, "12  常见问题 FAQ", level=1)
    faqs = [
        ("Q1. 网站打开显示 404 Not Found，怎么办？",
         "首次部署或刚部署完 10 分钟内 404 是正常的，等一会儿再刷新。如果很久了还是 404，打开仓库 → Settings → Pages，确认 Source 是「Deploy from a branch」，Branch 选「gh-pages」/「/(root)」，然后 Save。再等 1-2 分钟。"),
        ("Q2. 我添加了新素材，但网站上看不到？",
         "按顺序检查：1）后台是不是点了「保存草稿」；2）是不是执行了「导出全部数据」并覆盖 src/data/materials.json；3）封面图、预览视频是不是放进了 public/covers/ 和 public/previews/，且文件名完全一致；4）是不是执行了 ./deploy.sh；5）部署成功后有没有等 1-3 分钟；6）浏览器强制刷新（Cmd/Ctrl+Shift+R）。"),
        ("Q3. 预览视频在网页里打不开？",
         "检查格式：必须是 H.264 编码的 mp4（QuickTime 的 .mov、ProRes 都不支持网页直接播放）。如果是从夸克网盘直接取的直链，夸克可能不允许跨域播放 → 请用本地转码的低分辨率小样。"),
        ("Q4. 浏览器提示输入 GitHub 账号密码，应该填什么？",
         "用户名填 z1sj。密码不是你 GitHub 的登录密码，而是 ghp_ 开头的 Personal Access Token（PAT）。如果 Token 丢了，按附录 C 重新生成。"),
        ("Q5. ./deploy.sh 提示 permission denied？",
         "在终端执行一次：chmod +x deploy.sh，然后再运行。"),
        ("Q6. ./deploy.sh 提示 npm: command not found？",
         "本地 Node.js 没加到 PATH，把下面一行复制到终端执行一次再加 deploy.sh：\n"
         "export PATH=\"/Users/genhaosan/Library/Application Support/TRAE SOLO CN/ModularData/ai-agent/work-mode-projects/6a8bb8f8fe83435acf288eb9/.tools/node/bin:$PATH\""),
        ("Q7. 首页精选区有空白的卡片 / 数量不对？",
         "检查被你标为「首页精选」的素材数量是不是 0 或者太多（建议 6-10 个）。或者在站点设置里手动指定 featuredIds。"),
        ("Q8. 中文搜索搜不到结果？",
         "搜索会匹配「标题、描述、标签」三个字段。请确认关键词是否在这三个字段里出现。搜索是模糊匹配，错别字也有一定概率命中，但越准确越好。"),
        ("Q9. 上传大体积预览视频会影响吗？",
         "单个预览视频 > 30MB 会显著拖慢首屏加载。强烈建议用 FFmpeg 或剪映先转成 720P 低码率 H.264 版本。"),
        ("Q10. 我想把网站改成我自己的域名（比如 xxx.com）？",
         "在仓库 → Settings → Pages → Custom domain 里填入你的域名，然后去你的域名 DNS 服务商添加 CNAME 记录指向 z1sj.github.io。具体教程搜「GitHub Pages 自定义域名」有大量图文教程。"),
        ("Q11. 换电脑了，怎么继续更新？",
         "新电脑上安装 Git + Node.js（20 LTS 以上）→ 打开终端执行：git clone https://github.com/z1sj/third-frame-co.git → cd third-frame-co → npm install → 按本教程第 03 章流程开始操作即可。"),
    ]
    for q, a in faqs:
        p1 = doc.add_paragraph()
        p1.paragraph_format.left_indent = Cm(0.2)
        p1.paragraph_format.space_before = Pt(6)
        r1 = p1.add_run(q)
        set_run_font(r1, size=11, bold=True, color=RGBColor(0x00, 0x71, 0xE5))
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.6)
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(a)
        set_run_font(r2, size=10.5, color=RGBColor(0x3A, 0x3A, 0x3C))

    # ---------------------
    # 附录 A 字段速查表
    # ---------------------
    add_heading(doc, "附录 A  素材字段速查表", level=1)
    add_table(doc,
              ["字段 key", "类型 / 格式", "是否必填", "示例值"],
              [
                  ["id", "字符串（英文+数字+下划线）", "✅ 必填", "SONY_20260824_01"],
                  ["title", "字符串（≤30字推荐）", "✅ 必填", "逆光海边日落 S-Log3"],
                  ["type", "video / photo", "✅ 必填", "video"],
                  ["categoryPath", "字符串数组（3 个 ID）", "✅ 必填", "[\"camera\",\"sony\",\"slog3\"]"],
                  ["resolution", "4K / 1080P / 自定义", "推荐填", "4K"],
                  ["framerate", "整数", "推荐填", "25"],
                  ["scene", "indoor / outdoor / any", "推荐填", "outdoor"],
                  ["duration", "\"mm:ss\" 字符串", "视频必填", "00:42"],
                  ["quarkUrl", "URL 字符串", "✅ 必填", "https://pan.quark.cn/s/xxxxxx"],
                  ["baiduUrl", "URL 字符串", "可选", "https://pan.baidu.com/s/xxxxxx"],
                  ["coverFile", "文件名（含扩展名）", "✅ 必填", "SONY_20260824_01.jpg"],
                  ["previewFile", "文件名（含扩展名）", "视频必填", "SONY_20260824_01.mp4"],
                  ["description", "长文本", "✅ 必填", "…"],
                  ["tags", "字符串数组", "推荐填", "[\"逆光\",\"海边\",\"日落\"]"],
                  ["camera", "对象（model/lens/…）", "可选", "见类型定义"],
                  ["createdAt", "ISO 日期字符串", "自动", "2026-08-24T11:30:00Z"],
                  ["downloads", "整数", "自动", "0"],
                  ["views", "整数", "自动", "0"],
                  ["featured", "布尔值 true/false", "可选", "true"],
              ],
              col_widths_cm=[3.0, 4.2, 1.8, 7.0], first_col_bold=True)

    # ---------------------
    # 附录 B 目录结构
    # ---------------------
    add_heading(doc, "附录 B  项目目录结构", level=1)
    add_paragraph(doc, "了解目录结构有助于你定位文件：")
    add_code_block(doc,
                   "third-frame-co/\n"
                   "├── deploy.sh                      # 一键部署脚本（最常用）\n"
                   "├── package.json                   # Node 包依赖\n"
                   "├── vite.config.ts                 # Vite 构建配置（含 base 路径）\n"
                   "├── tailwind.config.js             # 配色、阴影、圆角等设计令牌\n"
                   "├── public/\n"
                   "│   ├── covers/                    # ⭐ 封面图目录，ID.jpg\n"
                   "│   └── previews/                  # ⭐ 预览视频目录，ID.mp4\n"
                   "└── src/\n"
                   "    ├── main.tsx                   # 入口文件\n"
                   "    ├── App.tsx                    # 路由\n"
                   "    ├── config.ts                  # ⭐ 密码、站点名、每页数量等配置\n"
                   "    ├── types/index.ts             # 所有 TypeScript 类型定义\n"
                   "    ├── data/\n"
                   "    │   ├── materials.json         # ⭐ 素材数据（后台导出后覆盖）\n"
                   "    │   ├── categories.json        # ⭐ 分类数据（后台导出后覆盖）\n"
                   "    │   ├── about.json             # 关于我数据\n"
                   "    │   └── demo.ts                # 演示备份数据\n"
                   "    ├── contexts/DataContext.tsx   # 全局数据、筛选、模糊搜索逻辑\n"
                   "    ├── components/                # UI 组件（卡片、过滤器、按钮…）\n"
                   "    └── pages/\n"
                   "        ├── Home.tsx               # 首页\n"
                   "        ├── Materials.tsx          # 素材列表 + 筛选 + 搜索\n"
                   "        ├── Detail.tsx             # 素材详情页\n"
                   "        ├── About.tsx              # 关于我\n"
                   "        └── Admin.tsx              # 后台管理面板\n")

    # ---------------------
    # 附录 C 重新生成 PAT
    # ---------------------
    add_heading(doc, "附录 C  重新生成 GitHub Personal Access Token", level=1)
    add_paragraph(doc, "场景：Token 丢失 / 过期 / 换电脑了 / git push 提示认证失败。")
    add_number(doc, 1, "浏览器打开 GitHub 并登录 → 右上角头像 → Settings")
    add_number(doc, 2, "左侧菜单拉到底 → Developer settings → Personal access tokens → Tokens (classic)")
    add_number(doc, 3, "点右上角 Generate new token → Generate new token (classic)，如果要你再输密码就输")
    add_number(doc, 4, "Note（备注）填 Third Frame Deploy（随便写，让你知道这是干嘛的）")
    add_number(doc, 5, "Expiration 选 No expiration 或者 90 days")
    add_number(doc, 6, "Select scopes 权限勾选：✅ repo（所有子项自动全勾）、✅ workflow、✅ read:org（在 admin:org 分组下面）")
    add_number(doc, 7, "最下面点绿色 Generate token 按钮")
    add_number(doc, 8, "复制 ghp_ 开头的那一串字符，粘贴到备忘录保存。页面关了以后就再也看不到这串了，丢了只能重新生成。")

    add_separator(doc)
    add_paragraph(doc, "—— 教程完 ——", size=10, color=RGBColor(0x8E, 0x8E, 0x93), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, f"文档生成于 {datetime.now().strftime('%Y-%m-%d %H:%M')} · Third Frame Co v1.0", size=9,
                  color=RGBColor(0xAE, 0xAE, 0xB2), align=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


def main():
    print(f"[1/2] 正在生成 Word 文档...")
    doc = build_document()
    # 确保输出目录存在
    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)
    doc.save(OUTPUT_FILE)
    size_kb = os.path.getsize(OUTPUT_FILE) / 1024
    print(f"[2/2] ✅ 已生成：{OUTPUT_FILE}  ({size_kb:.1f} KB)")
    print(f"      在 Finder 中打开：open \"{OUTPUT_FILE}\"")


if __name__ == "__main__":
    main()
